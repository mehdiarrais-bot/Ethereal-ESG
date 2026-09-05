"""
Lettre de mission / proposition d'accompagnement ESG (Word).

Document commercial du consultant : contexte du prospect, constats issus
du pré-diagnostic (couverture des exigences, maturité), objectifs, phases de
mission dérivées de la feuille de route réelle, livrables, prérequis.
Aucune donnée inventée — les honoraires restent à compléter.
"""
import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models import ESGRequest, ESGScores
from i18n import L
from docx_generator import THEME_HEX, DOCX_STYLES, hex_to_rgb, shade_cell, shade_paragraph, add_hr
from models import AestheticTheme


def _h(doc, text, size, color_hex, bold=True, space_before=14, space_after=4):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = hex_to_rgb(color_hex)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def generate_proposal_docx(request: ESGRequest, scores: ESGScores) -> bytes:
    from content_generator import compliance_assessment, maturity_text, roadmap_12m
    en = request.language == "en"
    colors = THEME_HEX.get(request.aesthetic_theme, THEME_HEX[AestheticTheme.CORPORATE_BLUE])
    if getattr(request, "custom_colors", None):
        from branding import brand_docx_hex
        colors = brand_docx_hex(colors, request.custom_colors)
    TR = L(request.language)
    name = request.company.name
    year = request.company.reporting_year

    gaps = [g for g in compliance_assessment(request, scores) if g["status"] in ("no", "partial")]
    mat = maturity_text(request, scores)
    mat_lbl = TR.get("mat_" + mat.get("key", "structured"), "")
    phases = roadmap_12m(request, scores)

    doc = Document()
    doc.styles['Normal'].font.size = Pt(10.5)
    for s in doc.sections:
        s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

    # ── En-tête ────────────────────────────────────────────────────────────
    kick = doc.add_paragraph()
    kr = kick.add_run(("PROPOSAL FOR ESG SUPPORT" if en else "PROPOSITION D'ACCOMPAGNEMENT ESG"))
    kr.bold = True; kr.font.size = Pt(10); kr.font.color.rgb = hex_to_rgb(colors["accent"])
    _h(doc, ("ESG engagement letter — " if en else "Lettre de mission ESG — ") + name,
       22, colors["primary"], space_before=2)
    sub = doc.add_paragraph()
    sr = sub.add_run(f"{request.company.sector}  ·  {request.company.country}  ·  "
                     + (f"FY {year}" if en else f"Exercice {year}"))
    sr.font.size = Pt(10.5); sr.font.color.rgb = hex_to_rgb("7F8C8D")
    add_hr(doc, colors["accent"])

    # ── 1. Contexte & pré-diagnostic ───────────────────────────────────────
    _h(doc, "1. " + ("Context and preliminary assessment" if en else "Contexte et pré-diagnostic"),
       14, colors["primary"])
    if en:
        ctx = (f"A preliminary review of {name}'s ESG data yields an overall score of "
               f"{scores.total_esg_score:.0f}/100 (indicative internal scale), with an ESG maturity "
               f"assessed as {mat_lbl.lower()} ({mat['stage']}/5). In view of the CSRD/VSME requirements "
               f"applicable to the SME/mid-cap market, the review identifies "
               f"{len(gaps)} regulatory gap(s) to address.")
    else:
        ctx = (f"Une revue préliminaire des données ESG de {name} aboutit à un score global de "
               f"{scores.total_esg_score:.0f}/100 (échelle interne indicative), pour une maturité ESG "
               f"évaluée comme {mat_lbl.lower()} ({mat['stage']}/5). Au regard des exigences CSRD/VSME "
               f"applicables au marché PME/ETI, cette revue identifie "
               f"{len(gaps)} écart(s) réglementaire(s) à traiter.")
    doc.add_paragraph(ctx)

    if gaps:
        import gap_status as GS   # source unique : pas de table locale
        tbl = doc.add_table(rows=len(gaps) + 1, cols=3)
        tbl.style = "Table Grid"
        heads = (TR["gap_req"], TR["gap_status"], TR["gap_note"])
        for ci, h in enumerate(heads):
            c = tbl.rows[0].cells[ci]
            run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shade_cell(c, colors["primary"])
        for ri, g in enumerate(gaps, 1):
            for ci, v in enumerate((g["req"],
                                    GS.libelle(TR, g["status"], g["nature"]),
                                    g["note"])):
                run = tbl.rows[ri].cells[ci].paragraphs[0].add_run(v)
                run.font.size = Pt(9)
                if ci == 1:
                    run.bold = True
                    run.font.color.rgb = hex_to_rgb("E74C3C" if g["status"] == "no" else "D97706")
        doc.add_paragraph()

    # ── 2. Objectifs de la mission ─────────────────────────────────────────
    _h(doc, "2. " + ("Engagement objectives" if en else "Objectifs de la mission"), 14, colors["primary"])
    objs_fr = ["Fiabiliser et compléter le reporting extra-financier au regard des exigences CSRD/VSME",
               "Traiter en priorité les écarts réglementaires identifiés au pré-diagnostic",
               "Structurer la gouvernance ESG et le plan d'action à 12 mois",
               "Produire les livrables de restitution à destination de la direction et des parties prenantes"]
    objs_en = ["Strengthen and complete sustainability reporting against CSRD/VSME requirements",
               "Address the regulatory gaps identified in the preliminary assessment as a priority",
               "Structure ESG governance and the 12-month action plan",
               "Produce management- and stakeholder-ready deliverables"]
    for o in (objs_en if en else objs_fr):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(o).font.size = Pt(10)

    # ── 3. Déroulé de la mission (issu de la feuille de route réelle) ──────
    _h(doc, "3. " + ("Proposed workplan" if en else "Déroulé proposé"), 14, colors["primary"])
    for ph in phases:
        if not ph["actions"]:
            continue
        p = doc.add_paragraph()
        r = p.add_run(f'{ph["label"]} — {ph["sub"]}')
        r.bold = True; r.font.size = Pt(11); r.font.color.rgb = hex_to_rgb(colors["secondary"])
        for act in ph["actions"][:4]:
            b = doc.add_paragraph(style="List Bullet")
            b.add_run(act["title"]).font.size = Pt(10)

    # ── 4. Livrables ───────────────────────────────────────────────────────
    _h(doc, "4. " + ("Deliverables" if en else "Livrables"), 14, colors["primary"])
    dels_fr = ["Rapport ESG complet (PDF) : diagnostic scoré par pilier, couverture des exigences de reporting, plan d'action",
               "Présentation de direction (PowerPoint) prête pour comité",
               "Rapport rédigé (Word) modifiable par vos équipes",
               "Feuille de route 12 mois priorisée (quick wins, responsables, échéances)"]
    dels_en = ["Full ESG report (PDF): scored diagnosis by pillar, reporting requirement coverage, action plan",
               "Board-ready management presentation (PowerPoint)",
               "Editable written report (Word)",
               "Prioritised 12-month roadmap (quick wins, owners, due dates)"]
    for d in (dels_en if en else dels_fr):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(d).font.size = Pt(10)

    # ── 5. Prérequis données ───────────────────────────────────────────────
    _h(doc, "5. " + ("Data requirements" if en else "Prérequis — données à réunir"), 14, colors["primary"])
    doc.add_paragraph(
        ("The engagement relies on the data listed in the collection template (energy, emissions, "
         "workforce, governance). Missing datapoints identified above will be collected jointly "
         "during the first phase." if en else
         "La mission s'appuie sur les données du modèle de collecte (énergie, émissions, effectifs, "
         "gouvernance). Les points de données manquants identifiés ci-dessus seront collectés "
         "conjointement lors de la première phase."))

    # ── 6. Conditions ──────────────────────────────────────────────────────
    _h(doc, "6. " + ("Terms" if en else "Conditions"), 14, colors["primary"])
    terms = doc.add_paragraph()
    tr_ = terms.add_run(
        ("Duration: 12 months from signature.  ·  Fees: [to complete]  ·  "
         "Payment terms: [to complete]  ·  Confidentiality: all data remains on the client's "
         "systems; processing is performed locally, with no external transfer." if en else
         "Durée : 12 mois à compter de la signature.  ·  Honoraires : [à compléter]  ·  "
         "Modalités de règlement : [à compléter]  ·  Confidentialité : les données restent sur les "
         "systèmes du client ; le traitement est réalisé en local, sans transfert externe."))
    tr_.font.size = Pt(10)
    shade_paragraph(terms, colors.get("light", "F0F2F5"))

    # ── Signatures ─────────────────────────────────────────────────────────
    doc.add_paragraph()
    sig = doc.add_table(rows=2, cols=2)
    labels = (("For the consultant", "For " + name) if en else ("Pour le consultant", "Pour " + name))
    for ci, lab in enumerate(labels):
        r0 = sig.rows[0].cells[ci].paragraphs[0].add_run(lab)
        r0.bold = True; r0.font.size = Pt(10)
        r1 = sig.rows[1].cells[ci].paragraphs[0].add_run(
            "\n\n" + ("Date & signature:" if en else "Date et signature :"))
        r1.font.size = Pt(9); r1.font.color.rgb = hex_to_rgb("7F8C8D")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
