import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from models import ESGRequest, ESGScores, AestheticTheme
from i18n import L

THEME_HEX = {
    AestheticTheme.CORPORATE_BLUE: {
        "light": "EBF2FB", "primary": "1B3A6B", "secondary": "2E86C1", "accent": "F39C12",
        "env": "27AE60", "social": "2E86C1", "gov": "8E44AD",
    },
    AestheticTheme.GREEN_NATURE: {
        "light": "D5F5E3", "primary": "1A5C38", "secondary": "27AE60", "accent": "F1C40F",
        "env": "2ECC71", "social": "3498DB", "gov": "E67E22",
    },
    AestheticTheme.DARK_PREMIUM: {
        "light": "E8EBF0", "primary": "1A1F28", "secondary": "58A6FF", "accent": "F7C948",
        "env": "3FB950", "social": "58A6FF", "gov": "BC8CFF",
    },
    AestheticTheme.MINIMAL_WHITE: {
        "light": "F5F5F5", "primary": "212121", "secondary": "1E88E5", "accent": "FF6F00",
        "env": "43A047", "social": "1E88E5", "gov": "8E24AA",
    },
    AestheticTheme.SUNSET_TERRACOTTA: {
        "light": "FAE5D8", "primary": "9A3412", "secondary": "E76F51", "accent": "F4A261",
        "env": "2A9D8F", "social": "E76F51", "gov": "6D597A",
    },
    AestheticTheme.OCEAN_DEEP: {
        "light": "DCF1F5", "primary": "0F4C5C", "secondary": "277DA1", "accent": "00BFA6",
        "env": "43AA8B", "social": "277DA1", "gov": "577590",
    },
    AestheticTheme.ROYAL_PURPLE: {
        "light": "EDE6F7", "primary": "2B1055", "secondary": "5E35B1", "accent": "D4A017",
        "env": "2E9E62", "social": "4A5FC1", "gov": "8E24AA",
    },
}


# Design language per theme: fonts + heading treatment
DOCX_STYLES = {
    AestheticTheme.CORPORATE_BLUE: {
        "font": "Calibri", "heading": "plain", "uppercase": False, "h1_size": 22,
    },
    AestheticTheme.GREEN_NATURE: {
        "font": "Trebuchet MS", "heading": "shaded", "uppercase": False, "h1_size": 17,
    },
    AestheticTheme.DARK_PREMIUM: {
        "font": "Georgia", "heading": "plain", "uppercase": True, "h1_size": 19,
    },
    AestheticTheme.MINIMAL_WHITE: {
        "font": "Segoe UI", "heading": "plain", "uppercase": True, "h1_size": 14,
    },
    AestheticTheme.SUNSET_TERRACOTTA: {
        "font": "Cambria", "heading": "shaded", "uppercase": False, "h1_size": 17,
    },
    AestheticTheme.OCEAN_DEEP: {
        "font": "Segoe UI", "heading": "plain", "uppercase": False, "h1_size": 20,
    },
    AestheticTheme.ROYAL_PURPLE: {
        "font": "Georgia", "heading": "plain", "uppercase": True, "h1_size": 19,
    },
}


def hex_to_rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def shade_paragraph(p, color_hex):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)


def add_heading(doc, text, level, color_hex, size=None, style=None):
    style = style or DOCX_STYLES[AestheticTheme.CORPORATE_BLUE]
    if style["uppercase"]:
        text = text.upper()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = style["font"]
    sizes = {1: style["h1_size"], 2: max(12, style["h1_size"] - 5), 3: 12}
    run.font.size = Pt(size or sizes.get(level, 12))
    if style["heading"] == "shaded" and level == 1:
        # Green Nature: white text on colored band
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_paragraph(p, color_hex)
    else:
        run.font.color.rgb = hex_to_rgb(color_hex)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_hr(doc, color_hex):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_kpi_table(doc, kpi_list, colors):
    if not kpi_list:
        return
    rows = [kpi_list[i:i+3] for i in range(0, len(kpi_list), 3)]
    for row_data in rows:
        while len(row_data) < 3:
            row_data.append(('', ''))
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        for col_idx, (label, value) in enumerate(row_data):
            # Label row
            cell = table.rows[0].cells[col_idx]
            shade_cell(cell, colors.get("light", "F5F5F5"))
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(label)
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = hex_to_rgb(colors["secondary"])
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Value row
            cell2 = table.rows[1].cells[col_idx]
            cell2.paragraphs[0].clear()
            run2 = cell2.paragraphs[0].add_run(str(value))
            run2.font.size = Pt(16)
            run2.font.bold = True
            run2.font.color.rgb = hex_to_rgb(colors["primary"])
            cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()


def add_score_block(doc, label, score, color_hex):
    p = doc.add_paragraph()
    run = p.add_run(f"{label} : ")
    run.font.size = Pt(11)
    run.bold = True
    score_run = p.add_run(f"{score:.1f}/100")
    score_run.font.size = Pt(14)
    score_run.bold = True
    score_run.font.color.rgb = hex_to_rgb(color_hex)


def add_bullet_list(doc, items, icon, color_hex):
    for item in items:
        p = doc.add_paragraph()
        icon_run = p.add_run(f"{icon}  ")
        icon_run.font.color.rgb = hex_to_rgb(color_hex)
        icon_run.font.size = Pt(10)
        text_run = p.add_run(item)
        text_run.font.size = Pt(10)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(3)


def generate_word_report(request: ESGRequest, scores: ESGScores, content: dict,
                         logo_bytes: bytes = None, cover_art: bytes = None,
                         charts: dict = None) -> bytes:
    colors = THEME_HEX.get(request.aesthetic_theme, THEME_HEX[AestheticTheme.CORPORATE_BLUE])
    style = DOCX_STYLES.get(request.aesthetic_theme, DOCX_STYLES[AestheticTheme.CORPORATE_BLUE])
    TR = L(request.language)

    from content_generator import pillar_headline, section_headlines
    _hl = pillar_headline(request, scores)
    _shl = section_headlines(request, scores)

    def add_conclusion(text, color_hex):
        """Sous-titre en gras = la conclusion de la section (information scent)."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12.5)
        run.font.color.rgb = hex_to_rgb(color_hex)
        p.paragraph_format.space_after = Pt(6)

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = style["font"]
    normal.font.size = Pt(10.5)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Cover ─────────────────────────────────────────────────────────────
    # Logo entreprise
    if logo_bytes:
        try:
            logo_p = doc.add_paragraph()
            logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            logo_p.add_run().add_picture(io.BytesIO(logo_bytes), height=Cm(2))
        except Exception:
            pass

    title_p = doc.add_paragraph()
    title_run = title_p.add_run(request.company.name.upper())
    title_run.font.size = Pt(32)
    title_run.bold = True
    title_run.font.name = style["font"]
    if request.aesthetic_theme == AestheticTheme.DARK_PREMIUM:
        # Luxe: white serif title on a dark band
        title_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_paragraph(title_p, "0D1117")
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif request.aesthetic_theme == AestheticTheme.GREEN_NATURE:
        title_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_paragraph(title_p, colors["primary"])
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        title_run.font.color.rgb = hex_to_rgb(colors["primary"])
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_after = Pt(4)

    add_hr(doc, colors["accent"])

    type_map = {
        "white_paper": TR["rep_white_paper"],
        "full_report": TR["rep_full_report"],
        "executive_summary_pdf": TR["rep_executive_summary_pdf"],
    }
    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(type_map.get(request.report_type.value, "Rapport ESG"))
    sub_run.font.size = Pt(18)
    sub_run.bold = True
    sub_run.font.color.rgb = hex_to_rgb(colors["secondary"])

    meta_p = doc.add_paragraph()
    meta_run = meta_p.add_run(
        f"{TR['exercise']} {request.company.reporting_year}  •  {request.company.sector}  •  {request.company.country}"
    )
    meta_run.font.size = Pt(11)
    meta_run.font.color.rgb = hex_to_rgb("7F8C8D")
    meta_p.paragraph_format.space_after = Pt(8)

    # Présentateur
    if request.company.presenter_name:
        pres_line = f"{TR['presented_by']} {request.company.presenter_name}"
        if request.company.presenter_title:
            pres_line += f" — {request.company.presenter_title}"
        pres_p = doc.add_paragraph()
        pres_run = pres_p.add_run(pres_line)
        pres_run.font.size = Pt(11)
        pres_run.italic = True
        pres_run.font.color.rgb = hex_to_rgb(colors["secondary"])
        pres_p.paragraph_format.space_after = Pt(12)

    # Illustration de couverture (générée localement)
    if cover_art:
        try:
            art_p = doc.add_paragraph()
            art_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            art_p.add_run().add_picture(io.BytesIO(cover_art), width=Cm(16))
            art_p.paragraph_format.space_after = Pt(12)
        except Exception:
            pass

    # Score summary table
    summary_table = doc.add_table(rows=2, cols=4)
    summary_table.style = 'Table Grid'
    headers = [TR["chart_env"], TR["chart_soc"], TR["chart_gov"], TR["score_global_short"]]
    values = [scores.environmental_score, scores.social_score,
              scores.governance_score, scores.total_esg_score]
    value_colors = [colors["env"], colors["social"], colors["gov"], colors["accent"]]

    for i, (hdr, val, col) in enumerate(zip(headers, values, value_colors)):
        hdr_cell = summary_table.rows[0].cells[i]
        shade_cell(hdr_cell, colors.get("light", "F5F5F5"))
        hdr_cell.paragraphs[0].clear()
        r = hdr_cell.paragraphs[0].add_run(hdr)
        r.font.size = Pt(9)
        r.bold = True
        r.font.color.rgb = hex_to_rgb(colors["secondary"])
        hdr_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        val_cell = summary_table.rows[1].cells[i]
        val_cell.paragraphs[0].clear()
        rv = val_cell.paragraphs[0].add_run(f"{val:.1f}")
        rv.font.size = Pt(24)
        rv.bold = True
        rv.font.color.rgb = hex_to_rgb(col)
        val_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    rating_p = doc.add_paragraph()
    rating_run = rating_p.add_run(f"{TR['note']} : {scores.rating}")
    rating_run.font.size = Pt(16)
    rating_run.bold = True
    rating_run.font.color.rgb = hex_to_rgb(colors["accent"])
    rating_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rating_p.paragraph_format.space_after = Pt(24)

    doc.add_page_break()

    # ── 1. Synthèse Exécutive ─────────────────────────────────────────────
    add_heading(doc, TR["pdf_s1"], 1, colors["primary"], style=style)
    add_hr(doc, colors["secondary"])
    exec_text = content.get("executive_summary",
        f"{request.company.name} présente son rapport ESG {request.company.reporting_year} "
        f"avec un score global de {scores.total_esg_score}/100 (note {scores.rating}).")
    doc.add_paragraph(exec_text).paragraph_format.space_after = Pt(12)

    # ── 2. Environnement ──────────────────────────────────────────────────
    add_heading(doc, TR["pdf_s2"], 1, colors["env"], style=style)
    add_hr(doc, colors["env"])
    add_conclusion(_hl["env"], colors["env"])
    add_score_block(doc, TR["score_env_label"], scores.environmental_score, colors["env"])

    env_text = content.get("environmental", "Analyse des données environnementales.")
    doc.add_paragraph(env_text).paragraph_format.space_after = Pt(8)

    env = request.environmental
    env_kpis = []
    if env.co2_emissions_tonnes is not None: env_kpis.append((TR["kpit"]["co2"], f"{env.co2_emissions_tonnes:,.0f}"))
    if env.renewable_energy_percent is not None: env_kpis.append((TR["kpit"]["renewable"], f"{env.renewable_energy_percent:.1f}%"))
    if env.energy_consumption_mwh is not None: env_kpis.append((TR["kpit"]["energy"], f"{env.energy_consumption_mwh:,.0f}"))
    if env.water_consumption_m3 is not None: env_kpis.append((TR["kpit"]["water"], f"{env.water_consumption_m3:,.0f}"))
    if env.waste_recycled_percent is not None: env_kpis.append((TR["kpit"]["recycling"], f"{env.waste_recycled_percent:.1f}%"))
    if env.scope1_emissions is not None: env_kpis.append((TR["kpit"]["s1"], f"{env.scope1_emissions:,.0f}"))
    if env.scope2_emissions is not None: env_kpis.append((TR["kpit"]["s2"], f"{env.scope2_emissions:,.0f}"))
    if env.scope3_emissions is not None: env_kpis.append((TR["kpit"]["s3"], f"{env.scope3_emissions:,.0f}"))
    add_kpi_table(doc, env_kpis, colors)

    # ── 3. Social ──────────────────────────────────────────────────────────
    add_heading(doc, TR["pdf_s3"], 1, colors["social"], style=style)
    add_hr(doc, colors["social"])
    add_conclusion(_hl["social"], colors["social"])
    add_score_block(doc, TR["score_soc_label"], scores.social_score, colors["social"])

    soc_text = content.get("social", "Analyse des données sociales.")
    doc.add_paragraph(soc_text).paragraph_format.space_after = Pt(8)

    soc = request.social
    soc_kpis = []
    if soc.total_employees is not None: soc_kpis.append((TR["kpit"]["employees"], f"{soc.total_employees:,}"))
    if soc.female_employees_percent is not None: soc_kpis.append((TR["kpit"]["women"], f"{soc.female_employees_percent:.1f}%"))
    if soc.employee_turnover_percent is not None: soc_kpis.append((TR["kpit"]["turnover"], f"{soc.employee_turnover_percent:.1f}%"))
    if soc.training_hours_per_employee is not None: soc_kpis.append((TR["kpit"]["training"], f"{soc.training_hours_per_employee:.0f}"))
    if soc.accident_frequency_rate is not None: soc_kpis.append((TR["kpit"]["accident"], f"{soc.accident_frequency_rate:.2f}"))
    if soc.customer_satisfaction_score is not None: soc_kpis.append((TR["kpit"]["satisfaction"], f"{soc.customer_satisfaction_score:.1f}"))
    add_kpi_table(doc, soc_kpis, colors)

    # ── 4. Gouvernance ────────────────────────────────────────────────────
    add_heading(doc, TR["pdf_s4"], 1, colors["gov"], style=style)
    add_hr(doc, colors["gov"])
    add_conclusion(_hl["gov"], colors["gov"])
    add_score_block(doc, TR["score_gov_label"], scores.governance_score, colors["gov"])

    gov_text = content.get("governance", "Analyse des données de gouvernance.")
    doc.add_paragraph(gov_text).paragraph_format.space_after = Pt(8)

    gov = request.governance
    gov_kpis = []
    if gov.board_members is not None: gov_kpis.append((TR["kpit"]["board"], str(gov.board_members)))
    if gov.female_board_percent is not None: gov_kpis.append((TR["kpit"]["women_board"], f"{gov.female_board_percent:.1f}%"))
    if gov.independent_board_percent is not None: gov_kpis.append((TR["kpit"]["independent"], f"{gov.independent_board_percent:.1f}%"))
    if gov.csr_budget_eur is not None: gov_kpis.append((TR["kpit"]["csr"], f"{gov.csr_budget_eur:,.0f}"))
    if gov.esg_audit_conducted is not None:
        gov_kpis.append((TR["kpit"]["audit"], TR["kpit"]["yes"] if gov.esg_audit_conducted else TR["kpit"]["no"]))
    if gov.sustainability_committee is not None:
        gov_kpis.append((TR["kpit"]["committee"], TR["kpit"]["yes"] if gov.sustainability_committee else TR["kpit"]["no"]))
    add_kpi_table(doc, gov_kpis, colors)

    doc.add_page_break()

    # ── 5. Analyses de Durabilité (CSRD / ESRS) ───────────────────────────
    charts = charts or {}

    def _docx_img(key, width_cm):
        if key in charts:
            try:
                pic_p = doc.add_paragraph()
                pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic_p.add_run().add_picture(io.BytesIO(charts[key]), width=Cm(width_cm))
                pic_p.paragraph_format.space_after = Pt(8)
            except Exception:
                pass

    add_heading(doc, TR["pdf_s5"], 1, colors["primary"], style=style)
    add_hr(doc, colors["secondary"])

    add_heading(doc, TR["sub_materiality"], 2, colors["secondary"], style=style)
    if content.get("materiality"):
        doc.add_paragraph(content["materiality"]).paragraph_format.space_after = Pt(6)
    _docx_img("materiality", 12)

    add_heading(doc, TR["sub_targets"], 2, colors["secondary"], style=style)
    if content.get("targets"):
        doc.add_paragraph(content["targets"]).paragraph_format.space_after = Pt(6)
    _docx_img("targets", 15)
    _docx_img("carbon_trajectory", 15)

    if content.get("taxonomy"):
        add_heading(doc, TR["sub_taxonomy"], 2, colors["secondary"], style=style)
        doc.add_paragraph(content["taxonomy"]).paragraph_format.space_after = Pt(6)
        _docx_img("taxonomy", 14)

    add_heading(doc, TR["sub_climate"], 2, colors["secondary"], style=style)
    if content.get("climate_risk"):
        doc.add_paragraph(content["climate_risk"]).paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # ── 6. Analyse Stratégique ────────────────────────────────────────────
    add_heading(doc, TR["pdf_s6"], 1, colors["primary"], style=style)
    add_hr(doc, colors["accent"])
    add_conclusion(_shl["strategic"], colors["accent"])

    add_heading(doc, TR["strengths"], 2, colors["env"], style=style)
    add_bullet_list(doc, scores.strengths, "✅", colors["env"])

    add_heading(doc, TR["weaknesses"], 2, "E74C3C", style=style)
    add_bullet_list(doc, scores.weaknesses, "⚠️", "E74C3C")

    if request.include_recommendations and scores.recommendations:
        from content_generator import enriched_recommendations
        add_heading(doc, TR["pdf_s7"], 1, colors["primary"], style=style)
        add_hr(doc, colors["accent"])
        for i, rec in enumerate(enriched_recommendations(request, scores), 1):
            p = doc.add_paragraph()
            num = p.add_run(f"{i}.  ")
            num.bold = True
            num.font.color.rgb = hex_to_rgb(colors["accent"])
            tr_ = p.add_run(rec["title"]); tr_.bold = True; tr_.font.size = Pt(10.5)
            p.paragraph_format.space_after = Pt(1)
            d = doc.add_paragraph()
            d.paragraph_format.left_indent = Cm(0.7)
            d.paragraph_format.space_after = Pt(6)
            dr = d.add_run(rec["detail"]); dr.font.size = Pt(9.5)
            meta = d.add_run(f'\n{TR["objective_col"]} : {rec["objective"]}  ·  '
                             f'{TR["owner_col"]} : {rec["owner"]}  ·  {rec["horizon"]}')
            meta.font.size = Pt(8.5)
            meta.font.color.rgb = hex_to_rgb("7F8C8D")

    # ── Diagnostic stratégique : benchmark + maturité + R&O + roadmap ─────
    from content_generator import benchmark_verdict, maturity_text, risks_opportunities, roadmap_12m
    _bv = benchmark_verdict(request, scores)
    _bm = _bv["bm"]
    _mt = maturity_text(request, scores)
    _ro = risks_opportunities(request, scores)

    add_heading(doc, TR["pdf_diag"], 1, colors["primary"], style=style)
    add_hr(doc, colors["accent"])
    p = doc.add_paragraph()
    r_ = p.add_run(_bv["title"]); r_.bold = True; r_.font.size = Pt(12)
    r_.font.color.rgb = hex_to_rgb(colors["secondary"])
    cap = doc.add_paragraph(TR["pdf_bench_sub"]); cap.runs[0].font.size = Pt(8)
    cap.runs[0].font.color.rgb = hex_to_rgb("7F8C8D")

    comp = {"env": scores.environmental_score, "social": scores.social_score,
            "gov": scores.governance_score, "global": scores.total_esg_score}
    rows = [(TR["bench_metric_col"], TR["bench_you"], TR["bench_sector"], TR["bench_delta_col"])]
    for key, lbl in [("env", TR["pillar_env"]), ("social", TR["pillar_soc"]),
                     ("gov", TR["pillar_gov"]), ("global", TR.get("score_global_short", "Global"))]:
        d = _bm["deltas"][key]
        rows.append((lbl, f"{comp[key]:.0f}", f"{_bm['avg'][key]:.0f}",
                     ("+" if d >= 0 else "") + f"{d:.0f}"))
    tbl = doc.add_table(rows=len(rows), cols=4)
    tbl.style = "Table Grid"
    for ci, val in enumerate(rows[0]):
        c = tbl.rows[0].cells[ci]; c.paragraphs[0].add_run(val).bold = True
        shade_cell(c, colors["primary"])
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].runs[0].font.size = Pt(9)
    for ri in range(1, len(rows)):
        for ci, val in enumerate(rows[ri]):
            c = tbl.rows[ri].cells[ci]; run = c.paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)
            if ci == 0:
                run.bold = True
            if ci == 3:
                run.bold = True
                run.font.color.rgb = hex_to_rgb("2E7D32" if not val.startswith("-") else "E74C3C")

    ins = doc.add_paragraph(); ins.paragraph_format.space_before = Pt(6)
    ins.add_run(_bv["insight"]).font.size = Pt(10)
    shade_paragraph(ins, colors.get("light", "F0F2F5"))

    _mat_lbl = TR.get("mat_" + _mt.get("key", "structured"), "")
    mp = doc.add_paragraph(); mp.paragraph_format.space_before = Pt(8)
    mr = mp.add_run(f'{TR["pdf_maturity_sub"]} : {_mat_lbl} ({_mt["stage"]}/5)')
    mr.bold = True; mr.font.color.rgb = hex_to_rgb(colors["secondary"]); mr.font.size = Pt(11)
    doc.add_paragraph(_mt["next_hint"])

    # ── Analyse des écarts réglementaires ─────────────────────────────────
    from content_generator import compliance_assessment
    _ga = compliance_assessment(request, scores)
    add_heading(doc, TR["gap_title"], 2, colors["secondary"], style=style)
    _st_hex = {"ok": "2E7D32", "partial": "D97706", "no": "E74C3C", "na": "7F8C8D"}
    _st_lbl = {"ok": TR["st_ok"], "partial": TR["st_partial"], "no": TR["st_no"], "na": TR["st_na"]}
    gtbl = doc.add_table(rows=len(_ga) + 1, cols=4)
    gtbl.style = "Table Grid"
    for ci, h in enumerate((TR["gap_req"], TR["gap_ref"], TR["gap_status"], TR["gap_note"])):
        c = gtbl.rows[0].cells[ci]
        run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(c, colors["primary"])
    for ri, g in enumerate(_ga, 1):
        vals = (g["req"], g["ref"], _st_lbl[g["status"]], g["note"])
        for ci, v in enumerate(vals):
            run = gtbl.rows[ri].cells[ci].paragraphs[0].add_run(v)
            run.font.size = Pt(9)
            if ci == 0:
                run.bold = True
            if ci == 2:
                run.bold = True
                run.font.color.rgb = hex_to_rgb(_st_hex[g["status"]])
    doc.add_paragraph()

    add_heading(doc, TR["risks_head"], 2, "E74C3C", style=style)
    add_bullet_list(doc, [
        f'[{i.get("priority", "P2")}] {i["tag"]} — {i["text"]} '
        f'({TR["risk_impact"].lower()} : {i.get("impact", "—").lower()} · '
        f'{TR["risk_lik"].lower()} : {i.get("likelihood", "—").lower()})'
        for i in _ro["risks"]], "▪", "E74C3C")
    add_heading(doc, TR["opps_head"], 2, colors["env"], style=style)
    add_bullet_list(doc, [f'{i["tag"]} — {i["text"]}' for i in _ro["opportunities"]], "▪", colors["env"])

    _rm = roadmap_12m(request, scores)
    if any(ph["actions"] for ph in _rm):
        add_heading(doc, TR["roadmap_title"], 2, colors["accent"], style=style)
        for ph in _rm:
            pp = doc.add_paragraph()
            hr = pp.add_run(f'{ph["label"]} — {ph["sub"]}')
            hr.bold = True; hr.font.color.rgb = hex_to_rgb(colors["primary"]); hr.font.size = Pt(10.5)
            for act in ph["actions"][:4]:
                ap = doc.add_paragraph(style="List Bullet")
                ar = ap.add_run(act["title"]); ar.font.size = Pt(10)
                if act["quick_win"]:
                    qw = ap.add_run(f'  [{TR["quick_win"]}]')
                    qw.bold = True; qw.font.size = Pt(8)
                    qw.font.color.rgb = hex_to_rgb(colors["accent"])

    # ── 7. Référentiels ───────────────────────────────────────────────────
    add_heading(doc, TR["pdf_s8"], 1, colors["primary"], style=style)
    add_hr(doc, colors["secondary"])
    ref_text = TR["ref_text_docx"]
    doc.add_paragraph(ref_text)

    # ── Conclusion ────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, TR["pdf_concl"], 1, colors["primary"], style=style)
    add_hr(doc, colors["primary"])
    conclusion = content.get("conclusion",
        f"{request.company.name} réaffirme son engagement vers un modèle d'affaires durable. "
        "Les axes d'amélioration identifiés feront l'objet de plans d'action concrets."
    )
    doc.add_paragraph(conclusion)

    # ── Note méthodologique ───────────────────────────────────────────────
    if content.get("methodology"):
        add_heading(doc, TR["pdf_methodo"], 2, colors["secondary"], style=style)
        mp = doc.add_paragraph()
        mr = mp.add_run(content["methodology"])
        mr.font.size = Pt(8.5)
        mr.font.color.rgb = hex_to_rgb("5A6572")
        shade_paragraph(mp, colors.get("light", "F0F2F5"))

    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        f"© {request.company.reporting_year} {request.company.name} — " + TR["gen_auto"]
    )
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = hex_to_rgb("7F8C8D")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
